#!/usr/bin/env python3
"""write_docx.py — generate .docx from JSON stdin payload.

Called by WriteOfficeTool (bundled PBS Python or ~/.deepseek/office-py venv).
Args: --output PATH   (output .docx path)
Stdin: JSON object with `title` (optional) and `blocks` array.

Blocks:
  { "type": "heading", "level": 1..6, "text": "..." }
  { "type": "paragraph", "text": "..." }
  { "type": "list", "style": "bullet"|"number", "items": ["..."] }
"""

import argparse
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"ERROR: {e}", file=sys.stderr)
    print("请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    payload = json.load(sys.stdin)
    doc = Document()

    # Title
    title = payload.get("title", "")
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Blocks
    blocks = payload.get("blocks", [])
    for block in blocks:
        typ = block.get("type", "paragraph")
        if typ == "heading":
            level = int(block.get("level", 1))
            level = max(1, min(level, 6))
            doc.add_heading(block.get("text", ""), level=level)
        elif typ == "paragraph":
            doc.add_paragraph(block.get("text", ""))
        elif typ == "list":
            style = block.get("style", "bullet")
            items = block.get("items", [])
            for item in items:
                if style == "number":
                    doc.add_paragraph(item, style="List Number")
                else:
                    doc.add_paragraph(item, style="List Bullet")

    doc.save(args.output)
    print("OK")


if __name__ == "__main__":
    main()
