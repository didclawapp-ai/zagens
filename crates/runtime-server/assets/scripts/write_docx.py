#!/usr/bin/env python3
"""write_docx.py — generate .docx from JSON stdin payload.

Called by WriteOfficeTool (bundled PBS Python or ~/.deepseek/office-py venv).
Args: --output PATH   (output .docx path)
Stdin: JSON object with `title`, `blocks`, and optional `page`, `header`, `footer`, `font`.

Blocks:
  { "type": "heading", "level": 1..6, "text": "..." }
  { "type": "paragraph", "text": "..." | "runs": [{ text, bold?, italic?, color?, size? }], "align"?: left|center|right|justify, "page_break_before"?: true }
  { "type": "list", "style": "bullet"|"number", "items": ["..."] | [{ text, subitems?: [...] }] }
  { "type": "table", "headers": [str], "rows": [[value...]], "style"?: str }
  { "type": "image", "path": "file.png", "width"?: px, "height"?: px, "caption"?: str }
  { "type": "toc", "title"?: "目录" }
"""

import argparse
import json
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError as e:
    print(f"ERROR: {e}", file=sys.stderr)
    print("请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAPER_SIZES = {
    "A4": (Cm(21.0), Cm(29.7)),
    "A3": (Cm(29.7), Cm(42.0)),
    "Letter": (Cm(21.59), Cm(27.94)),
}


def apply_runs(paragraph, runs):
    """Add formatted runs (rich text) to a paragraph."""
    for r in runs:
        text = str(r.get("text", ""))
        run = paragraph.add_run(text)
        if r.get("bold"):
            run.bold = True
        if r.get("italic"):
            run.italic = True
        color = r.get("color", "")
        if color and color.startswith("#") and len(color) == 7:
            run.font.color.rgb = RGBColor(
                int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            )
        size = r.get("size")
        if size is not None:
            run.font.size = Pt(float(size))


def set_page_setup(section, page_cfg):
    """Apply page settings to a section."""
    paper = page_cfg.get("paper", "A4")
    if paper in PAPER_SIZES:
        section.page_width, section.page_height = PAPER_SIZES[paper]

    orientation = page_cfg.get("orientation", "portrait")
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    margins = page_cfg.get("margins", {})
    if margins:
        section.top_margin = Cm(margins.get("top", 2.54))
        section.bottom_margin = Cm(margins.get("bottom", 2.54))
        section.left_margin = Cm(margins.get("left", 3.18))
        section.right_margin = Cm(margins.get("right", 3.18))


def set_header_footer(section, cfg):
    """Apply header/footer to a section."""
    if cfg.get("left") or cfg.get("center") or cfg.get("right"):
        header = section.header
        # Use a single paragraph with tab stops for left/center/right
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left = cfg.get("left", "")
        center = cfg.get("center", "")
        right = cfg.get("right", "")
        if left or center or right:
            p.clear()
            if left:
                p.add_run(left)
            if center:
                p.add_run("\t" + center)
            if right:
                p.add_run("\t" + right)
    elif cfg.get("text"):
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = cfg["text"]


def render_list(doc, block):
    """Render a list block, supporting nested subitems."""
    style = block.get("style", "bullet")
    style_name = "List Number" if style == "number" else "List Bullet"
    items = block.get("items", [])
    for item in items:
        if isinstance(item, dict):
            text = item.get("text", "")
            p = doc.add_paragraph(text, style=style_name)
            subitems = item.get("subitems", [])
            for si in subitems:
                si_text = si.get("text", si) if isinstance(si, dict) else str(si)
                doc.add_paragraph(si_text, style="List Bullet 2" if style == "bullet" else "List Number 2")
        else:
            doc.add_paragraph(str(item), style=style_name)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    payload = json.load(sys.stdin)
    doc = Document()

    # ── Page setup ──
    page_cfg = payload.get("page")
    if page_cfg:
        for section in doc.sections:
            set_page_setup(section, page_cfg)

    # ── Global font ──
    font_cfg = payload.get("font")
    if font_cfg:
        style = doc.styles["Normal"]
        font_obj = style.font
        if font_cfg.get("name"):
            font_obj.name = font_cfg["name"]
            # CJK font fallback
            rPr = style.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from lxml import etree as _ET
                rFonts = _ET.SubElement(rPr, qn("w:rFonts"))
            rFonts.set(qn("w:eastAsia"), font_cfg["name"])
        if font_cfg.get("size"):
            font_obj.size = Pt(float(font_cfg["size"]))

    # ── Title ──
    title = payload.get("title", "")
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Blocks ──
    blocks = payload.get("blocks") or []
    for block in blocks:
        typ = block.get("type", "paragraph")

        # -- heading --
        if typ == "heading":
            level = int(block.get("level", 1))
            level = max(1, min(level, 6))
            doc.add_heading(block.get("text", ""), level=level)

        # -- paragraph (plain text or rich text runs) --
        elif typ == "paragraph":
            # Page break before this paragraph?
            if block.get("page_break_before"):
                doc.add_page_break()

            runs = block.get("runs")
            if runs and isinstance(runs, list):
                p = doc.add_paragraph()
                apply_runs(p, runs)
            else:
                p = doc.add_paragraph(block.get("text", ""))

            # Alignment
            align = block.get("align")
            if align and align in ALIGN_MAP:
                p.alignment = ALIGN_MAP[align]

        # -- list --
        elif typ == "list":
            render_list(doc, block)

        # -- table --
        elif typ == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            table_style = block.get("style", "Light Grid Accent 1")
            if not headers and not rows:
                continue

            ncols = len(headers) if headers else (len(rows[0]) if rows else 1)
            nrows = len(rows) + (1 if headers else 0)
            table = doc.add_table(rows=nrows, cols=ncols)
            try:
                table.style = table_style
            except KeyError:
                pass  # table style not found, use default

            # Headers
            if headers:
                hdr_row = table.rows[0]
                for ci, h in enumerate(headers):
                    cell = hdr_row.cells[ci]
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                data_start = 1
            else:
                data_start = 0

            # Data rows
            for ri, row in enumerate(rows):
                r = table.rows[data_start + ri]
                for ci, val in enumerate(row):
                    if ci < ncols:
                        r.cells[ci].text = str(val) if val is not None else ""

        # -- image --
        elif typ == "image":
            img_path = block.get("path", "")
            if not img_path:
                continue

            # Resolve relative to the output file directory
            if not os.path.isabs(img_path):
                output_dir = os.path.dirname(os.path.abspath(args.output))
                img_path = os.path.join(output_dir, img_path)

            if os.path.isfile(img_path):
                width = block.get("width")
                height = block.get("height")
                kwargs = {}
                if width:
                    kwargs["width"] = Inches(float(width) / 96.0)
                if height:
                    kwargs["height"] = Inches(float(height) / 96.0)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_path, **kwargs)

                caption = block.get("caption", "")
                if caption:
                    cp = doc.add_paragraph(caption)
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.runs[0].italic = True if cp.runs else False

        # -- TOC --
        elif typ == "toc":
            toc_title = block.get("title", "目录")
            doc.add_heading(toc_title, level=1)
            # python-docx does not support real TOC field codes out of the box.
            # Insert a placeholder paragraph that instructs the user.
            doc.add_paragraph("（请在 Word 中右键此处 → 更新域以生成目录）")

    # ── Header / Footer ──
    header_cfg = payload.get("header")
    footer_cfg = payload.get("footer")
    if header_cfg or footer_cfg:
        for section in doc.sections:
            if header_cfg:
                set_header_footer(section, header_cfg)
            if footer_cfg:
                set_header_footer(section, footer_cfg)

    doc.save(args.output)
    print("OK")


if __name__ == "__main__":
    main()
